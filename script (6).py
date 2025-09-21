# Create utility modules and documentation

# logger.py - Logging utilities
logger_content = '''"""
Logger configuration for Cherry AI Assistant
"""

import logging
import logging.handlers
import os
from pathlib import Path
from typing import Dict, Any

def setup_logging(config: Dict[str, Any] = None) -> logging.Logger:
    """Setup logging configuration for Cherry"""
    
    if config is None:
        # Default configuration
        config = {
            'LOG_LEVEL': 'INFO',
            'LOG_FILE': 'data/logs/cherry.log',
            'MAX_LOG_FILES': 5,
            'LOGS_DIR': Path('data/logs')
        }
    
    # Create logs directory
    log_dir = Path(config.get('LOGS_DIR', 'data/logs'))
    log_dir.mkdir(parents=True, exist_ok=True)
    
    # Configure logging
    log_level = getattr(logging, config.get('LOG_LEVEL', 'INFO').upper())
    
    # Create formatter
    formatter = logging.Formatter(
        fmt='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
        datefmt='%Y-%m-%d %H:%M:%S'
    )
    
    # Configure root logger
    root_logger = logging.getLogger()
    root_logger.setLevel(log_level)
    
    # Clear existing handlers
    for handler in root_logger.handlers[:]:
        root_logger.removeHandler(handler)
    
    # Console handler
    console_handler = logging.StreamHandler()
    console_handler.setLevel(log_level)
    console_handler.setFormatter(formatter)
    root_logger.addHandler(console_handler)
    
    # File handler with rotation
    log_file = log_dir / 'cherry.log'
    file_handler = logging.handlers.RotatingFileHandler(
        log_file,
        maxBytes=10*1024*1024,  # 10MB
        backupCount=config.get('MAX_LOG_FILES', 5)
    )
    file_handler.setLevel(log_level)
    file_handler.setFormatter(formatter)
    root_logger.addHandler(file_handler)
    
    # Create Cherry logger
    cherry_logger = logging.getLogger('cherry')
    cherry_logger.info("Logging system initialized")
    
    return cherry_logger
'''

# file_handler.py - File operations
file_handler_content = '''"""
File Handler for Cherry AI Assistant
Handles file operations, document creation, and file management
"""

import os
import logging
import shutil
from pathlib import Path
from typing import Dict, Any, Optional, List
from datetime import datetime

try:
    from docx import Document
    import PyPDF2
except ImportError:
    print("File handling dependencies not installed. Install with:")
    print("pip install python-docx PyPDF2")

class FileHandler:
    """Handles file operations for Cherry"""
    
    def __init__(self, config: Dict[str, Any]):
        self.config = config
        self.logger = logging.getLogger(__name__)
        
        # Supported file types
        self.supported_text_formats = ['.txt', '.md', '.py', '.js', '.html', '.css', '.json']
        self.supported_doc_formats = ['.docx', '.pdf']
        
    async def create_file(self, filename: str, content: str, file_type: str = 'txt') -> Dict[str, Any]:
        """Create a new file with content"""
        try:
            # Ensure safe filename
            safe_filename = self._sanitize_filename(filename)
            
            # Add extension if not present
            if not Path(safe_filename).suffix:
                safe_filename += f'.{file_type}'
            
            # Create file path
            file_path = Path.cwd() / safe_filename
            
            # Write content based on file type
            if file_type.lower() == 'docx':
                success = await self._create_word_document(file_path, content)
            else:
                # Plain text file
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(content)
                success = True
            
            if success:
                self.logger.info(f"Created file: {file_path}")
                return {
                    'success': True,
                    'filepath': str(file_path),
                    'filename': safe_filename,
                    'size': file_path.stat().st_size
                }
            else:
                return {'success': False, 'error': 'Failed to create file'}
                
        except Exception as e:
            self.logger.error(f"Error creating file: {e}")
            return {'success': False, 'error': str(e)}
    
    async def read_file(self, filepath: str) -> Dict[str, Any]:
        """Read content from a file"""
        try:
            file_path = Path(filepath)
            
            if not file_path.exists():
                return {'success': False, 'error': 'File not found'}
            
            # Read content based on file type
            content = ""
            file_ext = file_path.suffix.lower()
            
            if file_ext == '.pdf':
                content = await self._read_pdf(file_path)
            elif file_ext == '.docx':
                content = await self._read_word_document(file_path)
            elif file_ext in self.supported_text_formats:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
            else:
                return {'success': False, 'error': f'Unsupported file type: {file_ext}'}
            
            return {
                'success': True,
                'content': content,
                'filepath': str(file_path),
                'size': file_path.stat().st_size,
                'modified': datetime.fromtimestamp(file_path.stat().st_mtime).isoformat()
            }
            
        except Exception as e:
            self.logger.error(f"Error reading file: {e}")
            return {'success': False, 'error': str(e)}
    
    async def list_files(self, directory: str = ".", pattern: str = "*") -> List[Dict[str, Any]]:
        """List files in a directory"""
        try:
            dir_path = Path(directory)
            
            if not dir_path.exists() or not dir_path.is_dir():
                return []
            
            files = []
            for file_path in dir_path.glob(pattern):
                if file_path.is_file():
                    stat = file_path.stat()
                    files.append({
                        'name': file_path.name,
                        'path': str(file_path),
                        'size': stat.st_size,
                        'modified': datetime.fromtimestamp(stat.st_mtime).isoformat(),
                        'extension': file_path.suffix
                    })
            
            return sorted(files, key=lambda x: x['modified'], reverse=True)
            
        except Exception as e:
            self.logger.error(f"Error listing files: {e}")
            return []
    
    async def delete_file(self, filepath: str) -> bool:
        """Delete a file"""
        try:
            file_path = Path(filepath)
            
            if file_path.exists():
                file_path.unlink()
                self.logger.info(f"Deleted file: {filepath}")
                return True
            else:
                self.logger.warning(f"File not found: {filepath}")
                return False
                
        except Exception as e:
            self.logger.error(f"Error deleting file: {e}")
            return False
    
    async def copy_file(self, source: str, destination: str) -> bool:
        """Copy a file"""
        try:
            source_path = Path(source)
            dest_path = Path(destination)
            
            if not source_path.exists():
                return False
            
            shutil.copy2(source_path, dest_path)
            self.logger.info(f"Copied file: {source} -> {destination}")
            return True
            
        except Exception as e:
            self.logger.error(f"Error copying file: {e}")
            return False
    
    async def move_file(self, source: str, destination: str) -> bool:
        """Move a file"""
        try:
            source_path = Path(source)
            dest_path = Path(destination)
            
            if not source_path.exists():
                return False
            
            shutil.move(str(source_path), str(dest_path))
            self.logger.info(f"Moved file: {source} -> {destination}")
            return True
            
        except Exception as e:
            self.logger.error(f"Error moving file: {e}")
            return False
    
    async def _create_word_document(self, filepath: Path, content: str) -> bool:
        """Create a Word document"""
        try:
            doc = Document()
            
            # Split content into paragraphs
            paragraphs = content.split('\\n\\n')
            
            for paragraph_text in paragraphs:
                if paragraph_text.strip():
                    doc.add_paragraph(paragraph_text.strip())
            
            doc.save(str(filepath))
            return True
            
        except Exception as e:
            self.logger.error(f"Error creating Word document: {e}")
            return False
    
    async def _read_word_document(self, filepath: Path) -> str:
        """Read content from a Word document"""
        try:
            doc = Document(str(filepath))
            content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text)
            
            return '\\n\\n'.join(content)
            
        except Exception as e:
            self.logger.error(f"Error reading Word document: {e}")
            return ""
    
    async def _read_pdf(self, filepath: Path) -> str:
        """Read content from a PDF file"""
        try:
            content = []
            
            with open(filepath, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page in pdf_reader.pages:
                    content.append(page.extract_text())
            
            return '\\n\\n'.join(content)
            
        except Exception as e:
            self.logger.error(f"Error reading PDF: {e}")
            return ""
    
    def _sanitize_filename(self, filename: str) -> str:
        """Sanitize filename to remove invalid characters"""
        import re
        
        # Remove invalid characters
        filename = re.sub(r'[<>:"/\\|?*]', '_', filename)
        
        # Remove leading/trailing spaces and dots
        filename = filename.strip('. ')
        
        # Limit length
        if len(filename) > 200:
            filename = filename[:200]
        
        # Ensure not empty
        if not filename:
            filename = f"file_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        
        return filename
    
    def get_supported_formats(self) -> Dict[str, List[str]]:
        """Get supported file formats"""
        return {
            'text': self.supported_text_formats,
            'document': self.supported_doc_formats
        }
'''

# web_scraper.py - Web information retrieval
web_scraper_content = '''"""
Web Scraper for Cherry AI Assistant
Handles web searches and information retrieval
"""

import asyncio
import logging
import aiohttp
from typing import Dict, Any, List, Optional
from urllib.parse import quote_plus
import json

try:
    from bs4 import BeautifulSoup
except ImportError:
    print("Web scraping dependencies not installed. Install with:")
    print("pip install beautifulsoup4 aiohttp")

class WebScraper:
    """Handles web searches and content extraction for Cherry"""
    
    def __init__(self, config: Dict[str, Any]):
        self.config = config
        self.logger = logging.getLogger(__name__)
        
        # Session for HTTP requests
        self.session = None
        
        # Search engines
        self.search_engines = {
            'duckduckgo': 'https://duckduckgo.com/html/?q={}',
            'bing': 'https://www.bing.com/search?q={}',
        }
        
        # Request headers
        self.headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
            'Accept-Language': 'en-US,en;q=0.5',
            'Accept-Encoding': 'gzip, deflate',
            'Connection': 'keep-alive',
            'Upgrade-Insecure-Requests': '1',
        }
    
    async def initialize(self):
        """Initialize the web scraper"""
        try:
            self.session = aiohttp.ClientSession(
                headers=self.headers,
                timeout=aiohttp.ClientTimeout(total=30)
            )
            self.logger.info("Web scraper initialized")
        except Exception as e:
            self.logger.error(f"Error initializing web scraper: {e}")
    
    async def search(self, query: str, max_results: int = 5) -> Dict[str, Any]:
        """Search for information on the web"""
        try:
            if not self.session:
                await self.initialize()
            
            self.logger.info(f"Searching web for: {query}")
            
            # Use DuckDuckGo for privacy-friendly search
            search_url = self.search_engines['duckduckgo'].format(quote_plus(query))
            
            async with self.session.get(search_url) as response:
                if response.status == 200:
                    html = await response.text()
                    results = await self._parse_duckduckgo_results(html, max_results)
                    
                    return {
                        'query': query,
                        'results': results,
                        'total_results': len(results)
                    }
                else:
                    self.logger.error(f"Search request failed: {response.status}")
                    return {'query': query, 'results': [], 'error': f'HTTP {response.status}'}
            
        except Exception as e:
            self.logger.error(f"Error searching web: {e}")
            return {'query': query, 'results': [], 'error': str(e)}
    
    async def _parse_duckduckgo_results(self, html: str, max_results: int) -> List[Dict[str, str]]:
        """Parse DuckDuckGo search results"""
        try:
            soup = BeautifulSoup(html, 'html.parser')
            results = []
            
            # Find result containers
            result_containers = soup.find_all('div', class_='result')
            
            for container in result_containers[:max_results]:
                try:
                    # Extract title
                    title_elem = container.find('a', class_='result__a')
                    title = title_elem.get_text().strip() if title_elem else "No title"
                    
                    # Extract URL
                    url = title_elem.get('href') if title_elem else ""
                    
                    # Extract snippet
                    snippet_elem = container.find('a', class_='result__snippet')
                    snippet = snippet_elem.get_text().strip() if snippet_elem else ""
                    
                    if title and url:
                        results.append({
                            'title': title,
                            'url': url,
                            'snippet': snippet
                        })
                        
                except Exception as e:
                    self.logger.warning(f"Error parsing result container: {e}")
                    continue
            
            return results
            
        except Exception as e:
            self.logger.error(f"Error parsing search results: {e}")
            return []
    
    async def get_page_content(self, url: str) -> Dict[str, Any]:
        """Get content from a web page"""
        try:
            if not self.session:
                await self.initialize()
            
            async with self.session.get(url) as response:
                if response.status == 200:
                    html = await response.text()
                    content = await self._extract_page_content(html)
                    
                    return {
                        'url': url,
                        'title': content.get('title', ''),
                        'text': content.get('text', ''),
                        'success': True
                    }
                else:
                    return {
                        'url': url,
                        'success': False,
                        'error': f'HTTP {response.status}'
                    }
        
        except Exception as e:
            self.logger.error(f"Error getting page content: {e}")
            return {
                'url': url,
                'success': False,
                'error': str(e)
            }
    
    async def _extract_page_content(self, html: str) -> Dict[str, str]:
        """Extract text content from HTML"""
        try:
            soup = BeautifulSoup(html, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Extract title
            title = ""
            title_tag = soup.find('title')
            if title_tag:
                title = title_tag.get_text().strip()
            
            # Extract main content
            # Try to find main content areas first
            main_content = soup.find('main') or soup.find('article') or soup.find('div', class_='content')
            
            if main_content:
                text = main_content.get_text()
            else:
                text = soup.get_text()
            
            # Clean up text
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = ' '.join(chunk for chunk in chunks if chunk)
            
            # Limit text length
            if len(text) > 5000:
                text = text[:5000] + "..."
            
            return {
                'title': title,
                'text': text
            }
            
        except Exception as e:
            self.logger.error(f"Error extracting page content: {e}")
            return {'title': '', 'text': ''}
    
    async def quick_search_summary(self, query: str) -> str:
        """Get a quick summary from web search"""
        try:
            search_results = await self.search(query, max_results=3)
            
            if search_results['results']:
                summary_parts = [f"Search results for '{query}':"]
                
                for i, result in enumerate(search_results['results'], 1):
                    summary_parts.append(f"{i}. {result['title']}")
                    if result['snippet']:
                        summary_parts.append(f"   {result['snippet']}")
                
                return '\\n'.join(summary_parts)
            else:
                return f"No search results found for '{query}'"
                
        except Exception as e:
            self.logger.error(f"Error getting search summary: {e}")
            return f"Error searching for '{query}': {str(e)}"
    
    async def cleanup(self):
        """Cleanup web scraper resources"""
        try:
            if self.session:
                await self.session.close()
            self.logger.info("Web scraper cleanup completed")
        except Exception as e:
            self.logger.error(f"Error during web scraper cleanup: {e}")
'''

# README.md - Project documentation
readme_content = '''# 🍒 Cherry AI Desktop Assistant

Cherry is an intelligent AI-powered desktop assistant that lives in your system tray and can control your computer, understand what's on your screen, and help with various tasks using voice or text commands.

## ✨ Features

### 🧠 AI Intelligence
- **Powered by Google Gemini API** (free tier compatible)
- **Contextual Memory**: Remembers conversations and learns your preferences
- **Multi-modal Understanding**: Processes text, voice, and visual input
- **RAG Integration**: Long-term memory with vector embeddings

### 🎤 Voice Capabilities
- **Wake Word Detection**: Say "Cherry" to activate
- **Speech Recognition**: Multiple engines (Whisper, Vosk, Google)
- **Text-to-Speech**: Natural voice responses with multiple voice options
- **Continuous Listening**: Background voice activation

### 👀 Computer Vision
- **Screen Analysis**: Understands what's currently on your screen
- **Screenshot Capture**: Take and analyze screenshots
- **OCR**: Extract text from images and screen content
- **Window Detection**: Identify and interact with applications

### ⌨️ Desktop Control
- **Mouse Control**: Click, drag, scroll at precise coordinates
- **Keyboard Automation**: Type text, press key combinations
- **Application Control**: Open, close, minimize, maximize windows
- **System Integration**: Native Windows, macOS, and Linux support

### 🖥️ Interface Options
- **System Tray**: Runs quietly in background
- **Popup GUI**: Modern Tkinter interface when needed
- **Global Hotkeys**: Customizable keyboard shortcuts
- **Voice Commands**: Hands-free operation

### 📁 File Operations
- **Document Creation**: Create Word docs, text files, code files
- **File Management**: Read, write, copy, move, delete files
- **Multi-format Support**: PDF, DOCX, TXT, MD, and more

### 🌐 Web Integration  
- **Web Search**: Search for information online
- **Content Extraction**: Read and summarize web pages
- **Privacy-Focused**: Uses DuckDuckGo for searches

## 🚀 Quick Start

### Prerequisites
- Python 3.8 or higher
- Windows 10/11, macOS 10.14+, or Linux
- Google Gemini API key (free at [aistudio.google.com](https://aistudio.google.com))

### Installation

1. **Clone or Download** the Cherry AI Assistant files
2. **Install Dependencies**:
   ```bash
   pip install -r requirements.txt
   ```

3. **Get Gemini API Key**:
   - Visit [Google AI Studio](https://aistudio.google.com)
   - Create a free account
   - Generate an API key

4. **Configure Cherry**:
   ```bash
   cp .env.example .env
   ```
   Edit `.env` and add your API key:
   ```
   GEMINI_API_KEY=your_api_key_here
   ```

5. **Run Cherry**:
   ```bash
   python main.py
   ```

Cherry will start in your system tray! 🍒

## 🎮 How to Use

### Voice Commands
1. Say **"Cherry"** to wake her up
2. Give your command: "Cherry, take a screenshot"
3. Cherry will respond and perform the action

### Keyboard Shortcuts
- `Ctrl+Alt+C`: Show Cherry interface
- `Ctrl+Alt+Q`: Quit Cherry
- `Escape`: Hide interface

### Text Interface
- Click the Cherry icon in system tray
- Select "Show Cherry" to open the GUI
- Type commands or chat naturally

## 🗣️ Example Commands

### Desktop Control
- "Cherry, take a screenshot"
- "Cherry, open calculator"
- "Cherry, type 'Hello World'"
- "Cherry, click at coordinates 100, 200"

### File Operations  
- "Cherry, create a document called 'meeting notes'"
- "Cherry, read the contents of readme.txt"
- "Cherry, list all Python files in this folder"

### Web & Research
- "Cherry, search for Python tutorials"
- "Cherry, what's the weather like today?"
- "Cherry, summarize this webpage"

### System Tasks
- "Cherry, what's running on my screen?"
- "Cherry, minimize all windows"
- "Cherry, show me system information"

## ⚙️ Configuration

### Environment Variables (.env)
```bash
# Required
GEMINI_API_KEY=your_api_key_here

# Optional Customization
WAKE_WORD=cherry
TTS_RATE=200
GUI_THEME=modern
LOG_LEVEL=INFO
```

### Voice Settings
- Change wake word in settings
- Adjust speech rate and volume
- Select different voice engines

### Memory Settings
- Control context window size
- Set memory retention period
- Configure RAG parameters

## 🏗️ Architecture

```
Cherry AI Assistant
├── Core Engine (cherry_brain.py)
│   ├── Gemini API Integration
│   ├── Task Planning & Execution
│   └── Context Management
├── Memory System (memory_manager.py)
│   ├── Vector Database (ChromaDB)
│   ├── Conversation History
│   └── RAG Implementation
├── Voice Processing (voice_processor.py)
│   ├── Wake Word Detection
│   ├── Speech Recognition
│   └── Text-to-Speech
├── Vision System (vision_system.py)
│   ├── Screen Capture
│   ├── OCR & Analysis
│   └── Window Detection
├── Desktop Control (desktop_controller.py)
│   ├── Mouse & Keyboard
│   ├── Application Control
│   └── System Integration
└── Interface Layer
    ├── System Tray (system_tray.py)
    ├── GUI Manager (gui_manager.py)
    └── Hotkey Manager (hotkey_manager.py)
```

## 🔧 Advanced Usage

### Custom Functions
Add your own capabilities by extending the Cherry brain:

```python
async def custom_function(self, params):
    # Your custom logic here
    return "Task completed!"
```

### Plugin System
Cherry's modular design makes it easy to add new features:
- Drop new modules in the `core/` directory
- Implement the standard async interface
- Register with the main brain

### API Integration
Connect Cherry to other services:
- Email (Gmail, Outlook)
- Calendar (Google Calendar)
- Cloud storage (Dropbox, Google Drive)
- Smart home devices

## 🛡️ Privacy & Security

- **Local Processing**: Most operations run locally
- **Encrypted Storage**: Sensitive data is encrypted
- **Privacy Mode**: Option to disable cloud features
- **Data Control**: Full control over stored information
- **No Tracking**: Cherry doesn't track or sell your data

## 📊 Free Tier Limits

Cherry is designed to work within Google Gemini's free tier:
- **15 requests per minute**
- **1,500 requests per day** 
- **1 million tokens per month**

Cherry optimizes usage by:
- Caching frequent responses
- Compressing context intelligently
- Using local processing when possible

## 🐛 Troubleshooting

### Common Issues

**"Cherry doesn't respond to wake word"**
- Check microphone permissions
- Adjust wake word sensitivity in settings
- Verify audio input device

**"GUI doesn't show"**
- Try `Ctrl+Alt+C` hotkey
- Check system tray for Cherry icon
- Restart Cherry if needed

**"API key errors"**
- Verify API key in `.env` file
- Check Gemini API quota
- Ensure API key has proper permissions

### Logs
Check `data/logs/cherry.log` for detailed error information.

## 🤝 Contributing

Cherry is open source! Contributions welcome:

1. Fork the repository
2. Create a feature branch
3. Add your improvements
4. Submit a pull request

### Development Setup
```bash
git clone <repository>
cd cherry-ai-assistant
pip install -r requirements.txt
pip install -r requirements-dev.txt  # Development dependencies
python -m pytest tests/  # Run tests
```

## 📄 License

This project is licensed under the MIT License - see the LICENSE file for details.

## 🙏 Acknowledgments

- Google Gemini for AI capabilities
- OpenAI for Whisper speech recognition
- The Python community for amazing libraries
- Contributors and testers

## 🔮 Roadmap

### Upcoming Features
- [ ] Multi-language support
- [ ] Calendar integration
- [ ] Email management
- [ ] Smart home control
- [ ] Custom voice training
- [ ] Mobile companion app
- [ ] Plugin marketplace
- [ ] Team collaboration features

### Performance Improvements
- [ ] Faster startup time
- [ ] Reduced memory usage
- [ ] Better caching strategies
- [ ] Offline mode enhancements

---

## 📞 Support

Need help? Here are your options:

1. **Documentation**: Check this README and code comments
2. **Issues**: Report bugs and request features on GitHub
3. **Discussions**: Join community discussions
4. **Wiki**: Detailed guides and tutorials

---

**Made with 🍒 and ❤️ for the AI community**

*Cherry AI Assistant - Your intelligent desktop companion*
'''

# Save all utility modules and documentation
with open('logger.py', 'w') as f:
    f.write(logger_content)

with open('file_handler.py', 'w') as f:
    f.write(file_handler_content)

with open('web_scraper.py', 'w') as f:
    f.write(web_scraper_content)

with open('README.md', 'w') as f:
    f.write(readme_content)

print("Created utility modules and documentation:")
print("✓ logger.py - Logging configuration")
print("✓ file_handler.py - File operations")
print("✓ web_scraper.py - Web information retrieval")
print("✓ README.md - Comprehensive project documentation")
print("\n🍒 Cherry AI Desktop Assistant Framework Complete! 🍒")
print("\nProject includes:")
print("• Complete AI brain with Gemini integration")
print("• Voice processing (speech recognition + TTS)")
print("• Computer vision (screen capture + analysis)")
print("• Desktop automation (mouse + keyboard control)")
print("• Memory management with RAG")
print("• System tray integration")
print("• GUI interface with Tkinter")
print("• Global hotkey support")
print("• File operations and web scraping")
print("• Comprehensive documentation")
print("• Production-ready requirements.txt")
print("\nReady to use with free Gemini API! 🚀")