"""
File Handler for Cherry AI Assistant
Handles file operations, document creation, and file management
"""

import os
import logging
import shutil
from pathlib import Path
from typing import Dict, Any, Optional, List
from datetime import datetime

try:
    from docx import Document
    import PyPDF2
except ImportError:
    print("File handling dependencies not installed. Install with:")
    print("pip install python-docx PyPDF2")

class FileHandler:
    """Handles file operations for Cherry"""

    def __init__(self, config: Dict[str, Any]):
        self.config = config
        self.logger = logging.getLogger(__name__)

        # Supported file types
        self.supported_text_formats = ['.txt', '.md', '.py', '.js', '.html', '.css', '.json']
        self.supported_doc_formats = ['.docx', '.pdf']

    async def create_file(self, filename: str, content: str, file_type: str = 'txt') -> Dict[str, Any]:
        """Create a new file with content"""
        try:
            # Ensure safe filename
            safe_filename = self._sanitize_filename(filename)

            # Add extension if not present
            if not Path(safe_filename).suffix:
                safe_filename += f'.{file_type}'

            # Create file path
            file_path = Path.cwd() / safe_filename

            # Write content based on file type
            if file_type.lower() == 'docx':
                success = await self._create_word_document(file_path, content)
            else:
                # Plain text file
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(content)
                success = True

            if success:
                self.logger.info(f"Created file: {file_path}")
                return {
                    'success': True,
                    'filepath': str(file_path),
                    'filename': safe_filename,
                    'size': file_path.stat().st_size
                }
            else:
                return {'success': False, 'error': 'Failed to create file'}

        except Exception as e:
            self.logger.error(f"Error creating file: {e}")
            return {'success': False, 'error': str(e)}

    async def read_file(self, filepath: str) -> Dict[str, Any]:
        """Read content from a file"""
        try:
            file_path = Path(filepath)

            if not file_path.exists():
                return {'success': False, 'error': 'File not found'}

            # Read content based on file type
            content = ""
            file_ext = file_path.suffix.lower()

            if file_ext == '.pdf':
                content = await self._read_pdf(file_path)
            elif file_ext == '.docx':
                content = await self._read_word_document(file_path)
            elif file_ext in self.supported_text_formats:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
            else:
                return {'success': False, 'error': f'Unsupported file type: {file_ext}'}

            return {
                'success': True,
                'content': content,
                'filepath': str(file_path),
                'size': file_path.stat().st_size,
                'modified': datetime.fromtimestamp(file_path.stat().st_mtime).isoformat()
            }

        except Exception as e:
            self.logger.error(f"Error reading file: {e}")
            return {'success': False, 'error': str(e)}

    async def list_files(self, directory: str = ".", pattern: str = "*") -> List[Dict[str, Any]]:
        """List files in a directory"""
        try:
            dir_path = Path(directory)

            if not dir_path.exists() or not dir_path.is_dir():
                return []

            files = []
            for file_path in dir_path.glob(pattern):
                if file_path.is_file():
                    stat = file_path.stat()
                    files.append({
                        'name': file_path.name,
                        'path': str(file_path),
                        'size': stat.st_size,
                        'modified': datetime.fromtimestamp(stat.st_mtime).isoformat(),
                        'extension': file_path.suffix
                    })

            return sorted(files, key=lambda x: x['modified'], reverse=True)

        except Exception as e:
            self.logger.error(f"Error listing files: {e}")
            return []

    async def delete_file(self, filepath: str) -> bool:
        """Delete a file"""
        try:
            file_path = Path(filepath)

            if file_path.exists():
                file_path.unlink()
                self.logger.info(f"Deleted file: {filepath}")
                return True
            else:
                self.logger.warning(f"File not found: {filepath}")
                return False

        except Exception as e:
            self.logger.error(f"Error deleting file: {e}")
            return False

    async def copy_file(self, source: str, destination: str) -> bool:
        """Copy a file"""
        try:
            source_path = Path(source)
            dest_path = Path(destination)

            if not source_path.exists():
                return False

            shutil.copy2(source_path, dest_path)
            self.logger.info(f"Copied file: {source} -> {destination}")
            return True

        except Exception as e:
            self.logger.error(f"Error copying file: {e}")
            return False

    async def move_file(self, source: str, destination: str) -> bool:
        """Move a file"""
        try:
            source_path = Path(source)
            dest_path = Path(destination)

            if not source_path.exists():
                return False

            shutil.move(str(source_path), str(dest_path))
            self.logger.info(f"Moved file: {source} -> {destination}")
            return True

        except Exception as e:
            self.logger.error(f"Error moving file: {e}")
            return False

    async def _create_word_document(self, filepath: Path, content: str) -> bool:
        """Create a Word document"""
        try:
            doc = Document()

            # Split content into paragraphs
            paragraphs = content.split('\n\n')

            for paragraph_text in paragraphs:
                if paragraph_text.strip():
                    doc.add_paragraph(paragraph_text.strip())

            doc.save(str(filepath))
            return True

        except Exception as e:
            self.logger.error(f"Error creating Word document: {e}")
            return False

    async def _read_word_document(self, filepath: Path) -> str:
        """Read content from a Word document"""
        try:
            doc = Document(str(filepath))
            content = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text)

            return '\n\n'.join(content)

        except Exception as e:
            self.logger.error(f"Error reading Word document: {e}")
            return ""

    async def _read_pdf(self, filepath: Path) -> str:
        """Read content from a PDF file"""
        try:
            content = []

            with open(filepath, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)

                for page in pdf_reader.pages:
                    content.append(page.extract_text())

            return '\n\n'.join(content)

        except Exception as e:
            self.logger.error(f"Error reading PDF: {e}")
            return ""

    def _sanitize_filename(self, filename: str) -> str:
        """Sanitize filename to remove invalid characters"""
        import re

        # Remove invalid characters
        filename = re.sub(r'[<>:"/\|?*]', '_', filename)

        # Remove leading/trailing spaces and dots
        filename = filename.strip('. ')

        # Limit length
        if len(filename) > 200:
            filename = filename[:200]

        # Ensure not empty
        if not filename:
            filename = f"file_{datetime.now().strftime('%Y%m%d_%H%M%S')}"

        return filename

    def get_supported_formats(self) -> Dict[str, List[str]]:
        """Get supported file formats"""
        return {
            'text': self.supported_text_formats,
            'document': self.supported_doc_formats
        }
